#!/usr/bin/python

# Standard libraries.
import os
import re
import sys
import shutil
import zipfile

# The python-docx library, for manipulating DOCX files. Importantly, when installing with pip,
# that's not the "docx" library, that's a different / earlier version - do "pip install python-docx"...
import docx
# ...and the docxcompose library for concatenating/appending DOCX files.
import docxcompose.composer
# ...and the docs-replace library, to install do "pip install python-docx-replace".
#import python_docx_replace

# We use Pandas to import Excel / CSV data.
#import numpy
import pandas

# Our own Docs To Markdown library.
import docsToMarkdownLib

# Extract the given DOCX file to a given temporary folder.
# Reads and returns the contents of the "document.xml" contained in the DOCX file.
def extractDocx(theFilename, destinationPath):
  templateDocx = zipfile.ZipFile(theFilename, "r")
  templateDocx.extractall(destinationPath)
  templateDocx.close()
  textHandle = open(destinationPath + "/word/document.xml")
  docxText = str(textHandle.read())
  textHandle.close()
  return(docxText)

# Turns the contents of the given folder into a DOCX file. Deletes the source folder when done.
def compressDocx(sourcePath, theFilename):
  theDocx = zipfile.ZipFile(theFilename, "w")
  for root, dirs, files in os.walk(sourcePath):
    for file in files:
      theDocx.write(os.path.join(root, file), os.path.join(root, file)[len(sourcePath):])
  theDocx.close()
  shutil.rmtree(sourcePath)

def putFile(thePath, theData):
  textHandle = open(thePath, "w")
  textHandle.write(theData)
  textHandle.close()

def processFolder(inputFolder, outputFolder):
  print("Processing Mailmerge folder: " + inputFolder + " to " + outputFolder, flush=True)
  
  # Go through items in the input folder, deal with particular items.
  synonyms = {}
  inputItems = []
  templateFiles = []
  defaultTemplate = "../default.docx"
  for inputItem in os.listdir(inputFolder):
    if os.path.isfile(inputFolder + os.sep + inputItem):
      # Get the filename and file type (extension, if there is one).
      fileSplit = inputItem.rsplit(".", 1)
      fileName = fileSplit[0]
      fileType = ""
      if len(fileSplit) > 1:
        fileType = fileSplit[1].upper()
        
      # Check for a "synonyms" XLSX (or XLS / CSV) file.
      if inputItem.lower() in["synonyms.xlsx", "synonyms.xls"]:
          for synonymIndex, synonymItem in pandas.read_excel(inputFolder + "/" + inputItem, header=None).iterrows():
            synonyms[synonymItem[0]] = synonymItem[1]
      elif inputItem.lower() in["synonyms.csv"]:
        for synonymIndex, synonymItem in pandas.read_csv(inputFolder + "/" + inputItem, header=None).iterrows():
            synonyms[synonymItem[0]] = synonymItem[1]
      # Otherwise, any XLSX / XLS / CSV file is an input file.
      elif fileType in ["XLSX", "XLS", "CSV"]:
        inputItems.append(inputItem)
      elif fileType in ["DOCX"]:
        templateFiles.append(inputItem)
  
  # Figure out the default template file - basically, any DOCX file that doesn't match an Excel file processed by this script.
  sortedTemplateFiles = []
  for inputItem in templateFiles:
    fileName = inputItem.rsplit(".", 1)[0]
    fileType = inputItem.rsplit(".", 1)[1].upper()
    if not inputItem in inputItems and not os.path.isdir(inputFolder + os.sep + fileName):
      sortedTemplateFiles.append(inputItem)
  sortedTemplateFiles = sorted(sortedTemplateFiles)
  if len(sortedTemplateFiles) > 0:
    defaultTemplate = sortedTemplateFiles[0]
  
  # Process each mailmerge data Excel (XLSX, XLS) or CSV file.
  foundHeadings = []
  for inputItem in inputItems:
    fileName = inputItem.rsplit(".", 1)[0]
    fileType = inputItem.rsplit(".", 1)[1].upper()
    mailData = pandas.DataFrame()
    mailArray = []
    if fileType in ["XLSX", "XLS"]:
      mailArray = pandas.read_excel(inputFolder + "/" + inputItem, sheet_name=None)
    elif fileType in ["CSV"]:
      mailArray = []
      mailArray.append(pandas.read_csv(inputFolder + "/" + inputItem))
    
    if not len(mailArray) == 0:
      print("Processing " + fileName + "...", flush=True)
      
      # Make sure there's an empty output folder with a name that matches the input filename.
      if os.path.isdir(outputFolder + os.sep + fileName):
        shutil.rmtree(outputFolder + os.sep + fileName)
      os.makedirs(outputFolder + os.sep + fileName, exist_ok=True)

      # Set up a Composer object to hold the multi-document version of the output.
      multiDoc = None #docx.Document()
      multiDocComposer = None #docxcompose.composer.Composer(multiDoc)
      
      for mailArraySheetName in mailArray.keys():
        outputPath = outputFolder + os.sep + fileName
        if len(mailArray) > 1:
          outputPath = outputPath + os.sep + mailArraySheetName
          os.makedirs(outputPath, exist_ok=True)
        mailData = mailArray[mailArraySheetName]
      
        # Make any column headers lower case for easier comparison.
        mailData.columns = map(str.lower, mailData.columns)
    
        # Process each item in the mailmerge data.
        for mailIndex, mailItem in mailData.iterrows():
          # Set the template file to use. See if there's a folder matching a column heading and template file matching a value, otherwise use the default.
          templateFile = defaultTemplate
          for heading in mailItem.index:
            heading = heading.lower()
            value = str(mailItem[heading]).lower()
            if value in synonyms.keys():
              value = synonyms[value].lower()
            if os.path.isdir(inputFolder + os.sep + heading):
              foundHeadings.append(heading)
            if os.path.isfile(inputFolder + os.sep + heading + os.sep + value + ".docx"):
              templateFile = inputFolder + os.sep + heading + os.sep + value + ".docx"

          if multiDoc == None:
            multiDoc = docx.Document()
            multiDocComposer = docxcompose.composer.Composer(multiDoc)

          # Open the template document as a plain text XML file.
          mailValues = mailItem.to_dict()
          docxText = extractDocx(inputFolder + os.sep + templateFile, "docxTemp")
          # Open the template document using python-docx.
          #mailDoc = docx.Document(inputFolder + os.sep + templateFile)

          # Note: this simple plain text search-and-replace seems to work okay, but in theory we might need to
          # be able to search for variable names (i.e. {{varname}}) split accross XML elements.
          
          # Extract a list of keys (i.e. variables to replace) from the document.
          mailKeys = []
          for docxMatch in re.finditer(r"\{\{.*?\}\}", docxText, re.MULTILINE | re.DOTALL):
            mailKeys.append(docxMatch.group(0)[2:-2])
          #mailKeys = python_docx_replace.docx_get_keys(mailDoc)

          # We skip any lines in the spreadsheet which have blank value for any variable included in
          # the template, otherwise we'll end up with an un-replaced variable string in the document.
          blankFound = False
          for mailKey in mailKeys:
            if mailKey.lower() in mailValues.keys():
              if pandas.isnull(mailValues[mailKey.lower()]) or mailValues[mailKey.lower()] in [None, ""]:
                blankFound = True
                
          if not blankFound:
            # Print a message for the user...
            print("Do Mailmerge: " + templateFile + " to " + outputPath + os.sep + str(mailIndex+1) + ".docx", flush=True)
            # ...replace key / value pairs...

            for mailKey in mailKeys:
              docxText = docxText.replace("{{" + mailKey + "}}", str(mailValues[mailKey]))
            #python_docx_replace.docx_replace(mailDoc, **mailValues)
          
            # ...and save the output document.
            #mailDoc.save(outputPath + os.sep + str(mailIndex+1) + ".docx")
            putFile("docxTemp/word/document.xml", docxText)
            compressDocx("docxTemp", outputPath + os.sep + str(mailIndex+1) + ".docx")

            mergedDoc = docx.Document(outputPath + os.sep + str(mailIndex+1) + ".docx")
            mergedDoc.add_page_break()
            multiDocComposer.append(mergedDoc)
    print("Writing: " + outputFolder + os.sep + fileName + ".docx")
    multiDocComposer.save(outputFolder + os.sep + fileName + ".docx")
        
  # Figure out if we want to resurse into any sub-folders.
  for inputItem in os.listdir(inputFolder):
    if os.path.isdir(inputFolder + os.sep + inputItem):
      if (not inputItem in inputItems) and (not inputItem.lower() in foundHeadings):
        processFolder(inputFolder + os.sep + inputItem, outputFolder + os.sep + inputItem)

processFolder(sys.argv[1], sys.argv[2])
